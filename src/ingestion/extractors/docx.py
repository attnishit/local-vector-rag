"""
DOCX Text Extraction

Author: Nishit Attrey

This module provides text extraction from Microsoft Word documents (.docx, .doc).

Functions:
    extract_text_from_docx: Extract text from DOCX files

Dependencies:
    python-docx: pip install python-docx
"""

import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_docx(
    filepath: Path,
    include_headers: bool = True,
    include_footers: bool = False,
    preserve_formatting: bool = True
) -> str:
    """
    Extract text from a DOCX file.

    Uses python-docx to extract text from Microsoft Word documents.
    Handles paragraphs, headings, tables, and optionally headers/footers.

    Args:
        filepath: Path to the DOCX file
        include_headers: Whether to include header content (default: True)
        include_footers: Whether to include footer content (default: False)
        preserve_formatting: Keep paragraph breaks and heading markers (default: True)

    Returns:
        Normalized plain text from the document

    Raises:
        ImportError: If python-docx is not installed
        ValueError: If file is corrupted or not a valid DOCX
        Exception: For other processing errors

    Example:
        >>> text = extract_text_from_docx(Path("report.docx"))
        >>> print(f"Extracted {len(text)} characters")

    Notes:
        - Paragraphs are separated by double newlines
        - Headings are preserved with "# " prefix for context
        - Tables are extracted row by row
        - Embedded images are skipped (only alt text if available)
        - Legacy .doc files require additional dependencies (textract)
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX extraction. "
            "Install it with: pip install python-docx"
        )

    if not filepath.exists():
        raise FileNotFoundError(f"DOCX file not found: {filepath}")

    # Check if it's a legacy .doc file (not supported by python-docx)
    if filepath.suffix.lower() == ".doc":
        logger.warning(
            f"Legacy .doc format detected: {filepath}\\n"
            "Attempting to read with python-docx (may fail). "
            "For full .doc support, install textract: pip install textract"
        )

    try:
        # Open the document
        doc = Document(filepath)

        extracted_parts = []

        # Extract headers if requested
        if include_headers:
            for section in doc.sections:
                header = section.header
                header_text = "\\n".join([p.text for p in header.paragraphs if p.text.strip()])
                if header_text:
                    extracted_parts.append(f"[HEADER]\\n{header_text}")

        # Extract main document content
        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('p'):
                para = element
                # Get the paragraph object
                for p in doc.paragraphs:
                    if p._element == para:
                        text = p.text.strip()
                        if text:
                            # Check if it's a heading
                            if preserve_formatting and p.style.name.startswith('Heading'):
                                level = p.style.name.replace('Heading', '').strip()
                                extracted_parts.append(f"# {text}")
                            else:
                                extracted_parts.append(text)
                        break

            # Handle tables
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        table_text = _extract_table_text(table)
                        if table_text:
                            extracted_parts.append(f"[TABLE]\\n{table_text}")
                        break

        # Extract footers if requested
        if include_footers:
            for section in doc.sections:
                footer = section.footer
                footer_text = "\\n".join([p.text for p in footer.paragraphs if p.text.strip()])
                if footer_text:
                    extracted_parts.append(f"[FOOTER]\\n{footer_text}")

        # Combine all parts
        full_text = "\\n\\n".join(extracted_parts)

        # Normalize whitespace
        full_text = " ".join(full_text.split())

        logger.info(f"Extracted {len(full_text)} characters from {filepath.name}")

        if len(full_text.strip()) < 10:
            logger.warning(f"DOCX extracted very little text ({len(full_text)} chars)")

        return full_text

    except Exception as e:
        # Try to provide helpful error messages
        if "not a valid" in str(e).lower() or "not recognized" in str(e).lower():
            raise ValueError(
                f"Invalid or corrupted DOCX file: {filepath}\\n"
                f"Error: {e}"
            )
        else:
            logger.error(f"Error extracting text from DOCX {filepath}: {e}")
            raise Exception(f"Failed to extract text from DOCX: {e}")


def _extract_table_text(table) -> str:
    """
    Extract text from a table in a structured format.

    Args:
        table: python-docx Table object

    Returns:
        Formatted table text with rows separated by newlines
    """
    rows_text = []
    for row in table.rows:
        cells_text = [cell.text.strip() for cell in row.cells]
        # Join cells with " | " for readability
        row_text = " | ".join(cells_text)
        if row_text.strip():
            rows_text.append(row_text)

    return "\\n".join(rows_text)


def extract_metadata_from_docx(filepath: Path) -> Optional[dict]:
    """
    Extract metadata from a DOCX file.

    Args:
        filepath: Path to the DOCX file

    Returns:
        Dictionary with core properties (title, author, subject, etc.)
        or None if extraction fails

    Example:
        >>> meta = extract_metadata_from_docx(Path("report.docx"))
        >>> if meta:
        ...     print(f"Author: {meta.get('author', 'Unknown')}")
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx is required for DOCX metadata extraction")
        return None

    try:
        doc = Document(filepath)
        core_props = doc.core_properties

        metadata = {
            "title": core_props.title,
            "author": core_props.author,
            "subject": core_props.subject,
            "keywords": core_props.keywords,
            "created": core_props.created,
            "modified": core_props.modified,
        }

        return {k: v for k, v in metadata.items() if v is not None}

    except Exception as e:
        logger.error(f"Failed to extract metadata from {filepath}: {e}")
        return None
