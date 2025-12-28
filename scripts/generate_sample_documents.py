#!/usr/bin/env python3
"""
Generate sample documents in PDF and DOCX formats for testing multi-format support.

This script creates sample documents that demonstrate the RAG system's ability
to handle various document formats.
"""

from pathlib import Path


def create_sample_pdf():
    """Create a sample PDF document using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("PyMuPDF not installed. Skipping PDF creation.")
        return False

    # Create a new PDF
    doc = fitz.open()

    # Page 1
    page1 = doc.new_page()
    text1 = """
Machine Learning Fundamentals

Machine learning is a subset of artificial intelligence that enables systems to learn and
improve from experience without being explicitly programmed. It focuses on developing
computer programs that can access data and use it to learn for themselves.

Types of Machine Learning:

1. Supervised Learning
   - Training data includes desired outputs
   - Examples: Classification, Regression
   - Algorithms: Linear Regression, Decision Trees, Neural Networks

2. Unsupervised Learning
   - Training data does not include outputs
   - Examples: Clustering, Dimensionality Reduction
   - Algorithms: K-Means, PCA, Autoencoders

3. Reinforcement Learning
   - Learning through interaction with environment
   - Reward-based learning
   - Examples: Game playing, Robotics
    """
    page1.insert_text((50, 50), text1, fontsize=11)

    # Page 2
    page2 = doc.new_page()
    text2 = """
Deep Learning and Neural Networks

Deep learning is a specialized branch of machine learning that uses neural networks
with multiple layers (hence "deep") to progressively extract higher-level features
from raw input.

Key Concepts:

- Neurons: Basic computational units
- Layers: Input layer, hidden layers, output layer
- Activation Functions: ReLU, Sigmoid, Tanh
- Backpropagation: Algorithm for training networks
- Gradient Descent: Optimization method

Common Architectures:
- Convolutional Neural Networks (CNNs) for image processing
- Recurrent Neural Networks (RNNs) for sequential data
- Transformers for natural language processing

Applications:
Computer vision, speech recognition, natural language processing, autonomous vehicles,
medical diagnosis, and many more domains benefit from deep learning techniques.
    """
    page2.insert_text((50, 50), text2, fontsize=11)

    # Save the PDF
    output_path = Path("data/raw/samples/sample_document.pdf")
    doc.save(output_path)
    doc.close()

    print(f"✓ Created PDF sample: {output_path}")
    return True


def create_sample_docx():
    """Create a sample DOCX document using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        print("python-docx not installed. Skipping DOCX creation.")
        return False

    # Create a new Document
    doc = Document()

    # Title
    title = doc.add_heading("Natural Language Processing Overview", level=1)

    # Introduction
    doc.add_paragraph(
        "Natural Language Processing (NLP) is a field of artificial intelligence "
        "that focuses on the interaction between computers and human language. "
        "It combines computational linguistics with machine learning and deep learning "
        "to enable computers to understand, interpret, and generate human language."
    )

    # Section 1
    doc.add_heading("Core NLP Tasks", level=2)

    doc.add_paragraph(
        "Text Classification: Assigning predefined categories to text documents. "
        "Examples include spam detection, sentiment analysis, and topic categorization."
    )

    doc.add_paragraph(
        "Named Entity Recognition (NER): Identifying and classifying named entities "
        "(persons, organizations, locations, dates) in text."
    )

    doc.add_paragraph(
        "Machine Translation: Automatically translating text from one language to another. "
        "Modern systems use neural machine translation (NMT) with attention mechanisms."
    )

    # Section 2
    doc.add_heading("Word Embeddings", level=2)

    doc.add_paragraph(
        "Word embeddings are dense vector representations of words that capture semantic relationships. "
        "Popular embedding methods include:"
    )

    # Bulleted list
    doc.add_paragraph("Word2Vec: Uses shallow neural networks to learn word associations", style="List Bullet")
    doc.add_paragraph("GloVe: Global Vectors for word representation based on co-occurrence statistics", style="List Bullet")
    doc.add_paragraph("FastText: Extension of Word2Vec that works with subword information", style="List Bullet")
    doc.add_paragraph("Contextual Embeddings: BERT, GPT, and other transformer-based models", style="List Bullet")

    # Section 3
    doc.add_heading("Modern NLP with Transformers", level=2)

    doc.add_paragraph(
        "The transformer architecture, introduced in 2017, revolutionized NLP. "
        "Key innovations include:"
    )

    # Add a simple table
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Model"
    header_cells[1].text = "Key Contribution"

    # Data rows
    table.rows[1].cells[0].text = "BERT"
    table.rows[1].cells[1].text = "Bidirectional encoding, masked language modeling"

    table.rows[2].cells[0].text = "GPT"
    table.rows[2].cells[1].text = "Autoregressive generation, zero-shot learning"

    table.rows[3].cells[0].text = "T5"
    table.rows[3].cells[1].text = "Text-to-text framework for all NLP tasks"

    # Conclusion
    doc.add_heading("Conclusion", level=2)

    doc.add_paragraph(
        "NLP continues to advance rapidly, with applications in chatbots, search engines, "
        "content generation, and information extraction. The field bridges the gap between "
        "human communication and machine understanding, opening new possibilities for "
        "human-computer interaction."
    )

    # Footer note
    doc.add_paragraph()
    footer = doc.add_paragraph(
        "This document was created to demonstrate multi-format support in the Local Vector RAG Database project."
    )
    footer.runs[0].italic = True

    # Save the document
    output_path = Path("data/raw/samples/sample_document.docx")
    doc.save(output_path)

    print(f"✓ Created DOCX sample: {output_path}")
    return True


def main():
    """Generate all sample documents."""
    print("Generating sample documents...")
    print()

    # Create samples directory if it doesn't exist
    Path("data/raw/samples").mkdir(parents=True, exist_ok=True)

    # Generate documents
    pdf_success = create_sample_pdf()
    docx_success = create_sample_docx()

    print()
    if pdf_success and docx_success:
        print("✓ All sample documents created successfully!")
        print()
        print("Sample files:")
        print("  - data/raw/samples/sample_document.pdf")
        print("  - data/raw/samples/sample_document.docx")
        print("  - data/raw/samples/sample_document.md")
        print()
        print("Test with:")
        print("  python main.py preview data/raw/samples/sample_document.pdf")
        print("  python main.py preview data/raw/samples/sample_document.docx")
        print("  python main.py preview data/raw/samples/sample_document.md")
    else:
        print("⚠ Some sample documents could not be created.")
        print("Ensure PyMuPDF and python-docx are installed:")
        print("  pip install PyMuPDF python-docx")


if __name__ == "__main__":
    main()
